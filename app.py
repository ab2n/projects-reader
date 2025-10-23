import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

st.title("Fiche projet Okoté")

# 1️⃣ Upload du fichier Excel
uploaded_file = st.file_uploader("Upload votre fichier Excel", type=["xlsx"])
if uploaded_file:
    df = pd.read_excel(uploaded_file)
    
    # 2️⃣ Choisir un projet
    projet_selectionne = st.selectbox("Sélectionnez un projet", df['Nom du Projet Okoté'])
    
    if projet_selectionne:
        projet_info = df[df['Nom du Projet Okoté'] == projet_selectionne].iloc[0]
        
        st.subheader("Informations du projet")
        for col in df.columns:
            st.write(f"**{col} :** {projet_info[col]}")
        
        # 3️⃣ Génération du DOCX
        def generate_docx(row):
            doc = Document()
            doc.add_heading(f"Fiche projet : {row['Nom du Projet Okoté']}", level=1)
            for col in df.columns:
                doc.add_paragraph(f"{col}: {row[col]}")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
        
        if st.button("Télécharger la fiche projet en DOCX"):
            docx_file = generate_docx(projet_info)
            st.download_button(
                label="Télécharger le DOCX",
                data=docx_file,
                file_name=f"Fiche_Projet_{projet_info['Nom du Projet Okoté']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
